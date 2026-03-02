from __future__ import annotations

import json
from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document

def _add_bullet(doc: Document, text: str) -> None:
    # Use built-in 'List Bullet' style if present; otherwise fall back to normal paragraph.
    try:
        p = doc.add_paragraph(text, style="List Bullet")
    except Exception:
        p = doc.add_paragraph(text)

def _add_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)

def _add_table(doc: Document, table_ref: str, caption: Optional[str] = None) -> None:
    # Minimal: table_ref placeholder. A real system would look up normalized tables by id.
    if caption:
        doc.add_paragraph(caption)
    t = doc.add_table(rows=1, cols=1)
    t.cell(0, 0).text = f"[table_ref: {table_ref}]"

def render_report_model_to_docx_bytes(report_model: Dict[str, Any], template_bytes: Optional[bytes] = None) -> bytes:
    """Render report_model-v1 into DOCX bytes using python-docx.

    Determinism considerations:
    - We avoid time-based metadata in the document.
    - Output is deterministic relative to template bytes + content ordering.
    """
    if template_bytes:
        doc = Document(BytesIO(template_bytes))
    else:
        doc = Document()

    report = report_model.get("report", {})
    title = report.get("title", "").strip()
    if title:
        doc.add_heading(title, level=0)

    # Summary
    summary = report.get("summary", [])
    if summary:
        doc.add_heading("Summary", level=1)
        for block in summary:
            btype = block.get("type")
            if btype == "bullet":
                _add_bullet(doc, block.get("text", ""))
            elif btype == "paragraph":
                _add_paragraph(doc, block.get("text", ""))
            elif btype == "table":
                _add_table(doc, block.get("table_ref", ""), block.get("caption"))
            else:
                _add_paragraph(doc, block.get("text", ""))

    # Sections
    sections = report.get("sections", [])
    for sec in sections:
        heading = sec.get("heading", "")
        if heading:
            doc.add_heading(heading, level=1)
        for block in sec.get("blocks", []):
            btype = block.get("type")
            if btype == "bullet":
                _add_bullet(doc, block.get("text", ""))
            elif btype == "paragraph":
                _add_paragraph(doc, block.get("text", ""))
            elif btype == "table":
                _add_table(doc, block.get("table_ref", ""), block.get("caption"))
            else:
                _add_paragraph(doc, block.get("text", ""))

    # Risks (optional)
    risks = report.get("risks", [])
    if risks:
        doc.add_heading("Risks", level=1)
        for block in risks:
            btype = block.get("type")
            if btype == "bullet":
                _add_bullet(doc, block.get("text", ""))
            else:
                _add_paragraph(doc, block.get("text", ""))

    # Save
    out = BytesIO()
    doc.save(out)
    return out.getvalue()
